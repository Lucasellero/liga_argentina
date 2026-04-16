"""
Análisis de Rebotes Ofensivos de Boca Juniors
Liga Nacional Argentina 2025-26
- OREB promedio, OREB%, Puntos de Segunda Oportunidad
- Split por Victoria / Derrota, en cada período
"""

import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

BOX = 'liga_argentina/docs/liga_nacional/liga_nacional.csv'
PBP = 'liga_argentina/docs/liga_nacional/liga_nacional_pbp.csv'

box = pd.read_csv(BOX)
pbp = pd.read_csv(PBP)

box['Fecha_dt'] = pd.to_datetime(box['Fecha'], format='%d/%m/%Y')
pbp['Fecha_dt'] = pd.to_datetime(pbp['Fecha'], format='%d/%m/%Y')

CUTOFF = pd.Timestamp('2026-02-02')

boca_totals = box[(box['Equipo'] == 'BOCA') & (box['Apellido'] == 'TOTALES')].copy()
game_dates  = sorted(boca_totals['Fecha_dt'].unique())
last5_dates = game_dates[-5:]

def get_dates(label):
    if label == 'PRE':  return [d for d in game_dates if d < CUTOFF]
    if label == 'CASA': return [d for d in game_dates if d >= CUTOFF]
    if label == 'L5':   return last5_dates

# ─── Map IdPartido → (Ganado, Fecha_dt) for Boca ────────────────────────────
game_info = boca_totals.set_index('IdPartido')[['Ganado', 'Fecha_dt']].copy()

# ─── BOCA PBP: only games where Boca played ─────────────────────────────────
boca_game_ids = set(boca_totals['IdPartido'].unique())

boca_pbp = pbp[pbp['IdPartido'].isin(boca_game_ids)].copy()
game_info_reset = game_info[['Ganado']].reset_index()  # IdPartido as column (Fecha_dt already in pbp)
boca_pbp = boca_pbp.merge(game_info_reset, on='IdPartido', how='left')

# ─── Determine which side is BOCA in each game ───────────────────────────────
# Equipo_lado = LOCAL or VISITANTE; we need to know if Boca is LOCAL or VISITANTE
boca_side_map = {}
for gid in boca_game_ids:
    sample = boca_pbp[boca_pbp['IdPartido'] == gid]
    # Find a row where Boca's player appears
    for _, r in sample.iterrows():
        if r['Equipo_lado'] in ('LOCAL', 'VISITANTE'):
            # Check box score to see which team name is Boca
            break
    # Better: look at Equipo_local and Equipo_visitante
    row = sample.iloc[0]
    if row['Equipo_local'] == 'BOCA':
        boca_side_map[gid] = 'LOCAL'
    else:
        boca_side_map[gid] = 'VISITANTE'

boca_pbp['boca_side'] = boca_pbp['IdPartido'].map(boca_side_map)
# A row's side is Boca's side if Equipo_lado == boca_side
boca_pbp['is_boca_action'] = boca_pbp['Equipo_lado'] == boca_pbp['boca_side']

# ─── SECOND-CHANCE POINTS via PBP ────────────────────────────────────────────
# For each game, walk through the PBP sequence:
#   1. Detect REBOTE-OFENSIVO by Boca
#   2. Track the ensuing possession until it ends (score, turnover, opp DREB, period end)
#   3. Sum points scored in that possession

POSSESSION_ENDERS = {
    'REBOTE-DEFENSIVO',  # opponent grabs → possession over
    'PERDIDA',           # Boca turnover → possession over
    'FINAL-PERIODO',     # period ends
    'FINAL-PARTIDO',
}

FG_EVENTS = {
    'CANASTA-2P', 'CANASTA-3P', 'TIRO2-FALLADO', 'TIRO3-FALLADO'
}

def pts_for_event(tipo):
    if tipo == 'CANASTA-3P': return 3
    if tipo == 'CANASTA-2P': return 2
    if tipo == 'CANASTA-1P': return 1
    return 0

second_chance_records = []  # (IdPartido, Ganado, Fecha_dt, pts_scored)

for gid, grp in boca_pbp.groupby('IdPartido'):
    grp = grp.sort_values('NumAccion').reset_index(drop=True)
    boca_side = boca_side_map.get(gid)
    ganado   = grp['Ganado'].iloc[0]
    fecha    = grp['Fecha_dt'].iloc[0]

    n = len(grp)
    i = 0
    while i < n:
        row = grp.iloc[i]
        # Detect Boca OREB
        if (row['Tipo'] == 'REBOTE-OFENSIVO' and
                row['Equipo_lado'] == boca_side):
            # Walk forward to track the ensuing possession
            pts = 0
            j = i + 1
            while j < n:
                nxt = grp.iloc[j]
                nxt_tipo = str(nxt['Tipo'])

                # Opponent DREB (possession ends, no points)
                if (nxt_tipo == 'REBOTE-DEFENSIVO' and
                        nxt['Equipo_lado'] != boca_side):
                    break

                # Boca DREB (shouldn't happen mid-possession, but end tracking)
                if (nxt_tipo == 'REBOTE-DEFENSIVO' and
                        nxt['Equipo_lado'] == boca_side):
                    break

                # Period / game end
                if nxt_tipo in ('FINAL-PERIODO', 'FINAL-PARTIDO'):
                    break

                # Boca turnover → possession ends
                if nxt_tipo == 'PERDIDA' and nxt['Equipo_lado'] == boca_side:
                    break

                # Boca scores (FT or FG by Boca)
                if nxt['Equipo_lado'] == boca_side and nxt_tipo.startswith('CANASTA'):
                    pts += pts_for_event(nxt_tipo)

                # After a made 2P or 3P the possession ends (opponent now has ball)
                # But FTs can continue — handle: if CANASTA-2P or CANASTA-3P → possession ends
                if (nxt_tipo in ('CANASTA-2P', 'CANASTA-3P') and
                        nxt['Equipo_lado'] == boca_side):
                    j += 1
                    break

                # Another Boca OREB (possession continues — already handled in outer loop
                # but we keep tracking from here)
                if (nxt_tipo == 'REBOTE-OFENSIVO' and
                        nxt['Equipo_lado'] == boca_side):
                    j += 1
                    continue

                # Boca missed shot (TIRO2/3-FALLADO) → possession may continue via OREB
                j += 1

            second_chance_records.append({
                'IdPartido': gid,
                'Ganado':    ganado,
                'Fecha_dt':  fecha,
                'pts':       pts,
            })
        i += 1

sc_df = pd.DataFrame(second_chance_records)

# ─── OREB from BOX SCORE (more reliable totals) ──────────────────────────────
# Also get opponent DREB for OREB%

def oreb_analysis(label):
    dates  = get_dates(label)
    boca_g = boca_totals[boca_totals['Fecha_dt'].isin(dates)].copy()
    boca_g = boca_g.set_index('IdPartido')

    # Opponent totals
    opp_rows = []
    for gid in boca_g.index:
        rival = boca_g.loc[gid, 'Rival']
        opp = box[(box['IdPartido'] == gid) &
                  (box['Equipo'] == rival) &
                  (box['Apellido'] == 'TOTALES')]
        opp_rows.append(opp)
    opp_df = pd.concat(opp_rows).set_index('IdPartido') if opp_rows else pd.DataFrame()

    results = []
    for won in [True, False]:
        mask = boca_g['Ganado'] == won
        sub  = boca_g[mask]
        n    = len(sub)
        if n == 0:
            continue

        boca_oreb = sub['OReb'].sum()
        boca_dreb = sub['DReb'].sum()

        opp_sub  = opp_df.loc[opp_df.index.isin(sub.index)]
        opp_dreb = opp_sub['DReb'].sum()
        opp_oreb = opp_sub['OReb'].sum()

        oreb_pct = boca_oreb / (boca_oreb + opp_dreb) * 100 if (boca_oreb + opp_dreb) > 0 else 0
        dreb_pct = boca_dreb / (boca_dreb + opp_oreb) * 100 if (boca_dreb + opp_oreb) > 0 else 0

        # Second-chance points for this subset
        sc_sub = sc_df[sc_df['IdPartido'].isin(sub.index) & (sc_df['Ganado'] == won)]
        total_sc_pts  = sc_sub['pts'].sum()
        oreb_count_pbp = len(sc_sub)  # one record per OREB event
        sc_ppp = total_sc_pts / oreb_count_pbp if oreb_count_pbp > 0 else 0

        results.append({
            'Escenario':        'Victoria' if won else 'Derrota',
            'Partidos':         n,
            'OREB_total':       int(boca_oreb),
            'OREB_pg':          round(boca_oreb / n, 1),
            'OREB_pct':         round(oreb_pct, 1),
            'SC_oportunidades': oreb_count_pbp,
            'SC_opp_pg':        round(oreb_count_pbp / n, 1),
            'SC_pts_total':     int(total_sc_pts),
            'SC_pts_pg':        round(total_sc_pts / n, 1),
            'SC_PPP':           round(sc_ppp, 3),
        })

    return results

# ─── PRINT RESULTS ───────────────────────────────────────────────────────────

PERIOD_LABELS = {
    'PRE':  'Pre-Casalanguida',
    'CASA': 'Era Casalanguida',
    'L5':   'Últimos 5 partidos',
}

COL_WIDTHS = {
    'Escenario': 10, 'Partidos': 8, 'OREB_total': 11, 'OREB_pg': 8,
    'OREB_pct': 8, 'SC_oportunidades': 16, 'SC_opp_pg': 10,
    'SC_pts_total': 13, 'SC_pts_pg': 10, 'SC_PPP': 7,
}

HEADERS = {
    'Escenario': 'Escenario', 'Partidos': 'PJ',
    'OREB_total': 'OREB Total', 'OREB_pg': 'OREB/G',
    'OREB_pct': 'OREB%', 'SC_oportunidades': 'SC Oportunidades',
    'SC_opp_pg': 'SC Opp/G', 'SC_pts_total': 'SC Pts Total',
    'SC_pts_pg': 'SC Pts/G', 'SC_PPP': 'PPP',
}

def print_table(rows):
    keys = list(COL_WIDTHS.keys())
    # Header
    header = '  '.join(HEADERS[k].ljust(COL_WIDTHS[k]) for k in keys)
    sep    = '  '.join('-' * COL_WIDTHS[k] for k in keys)
    print(header)
    print(sep)
    for r in rows:
        line = '  '.join(str(r[k]).ljust(COL_WIDTHS[k]) for k in keys)
        print(line)

for lbl, name in PERIOD_LABELS.items():
    n_games = len(get_dates(lbl))
    print(f'\n{"="*80}')
    print(f'  {name}  ({n_games} partidos)')
    print(f'{"="*80}')
    rows = oreb_analysis(lbl)
    if rows:
        print_table(rows)
        # Differential
        if len(rows) == 2:
            w = rows[0]
            l = rows[1]
            print(f'\n  Diferencial Victoria vs Derrota:')
            print(f'    OREB/G:   {w["OREB_pg"]:+.1f} → {l["OREB_pg"]:+.1f}  (Δ {w["OREB_pg"]-l["OREB_pg"]:+.1f})')
            print(f'    OREB%:    {w["OREB_pct"]:+.1f}% → {l["OREB_pct"]:+.1f}%  (Δ {w["OREB_pct"]-l["OREB_pct"]:+.1f}pp)')
            print(f'    SC Pts/G: {w["SC_pts_pg"]:+.1f} → {l["SC_pts_pg"]:+.1f}  (Δ {w["SC_pts_pg"]-l["SC_pts_pg"]:+.1f})')
            print(f'    PPP:      {w["SC_PPP"]:.3f} → {l["SC_PPP"]:.3f}  (Δ {w["SC_PPP"]-l["SC_PPP"]:+.3f})')
    else:
        print('  Sin datos')

print('\n')
print('='*80)
print('  NOTAS METODOLÓGICAS')
print('='*80)
print("""
  OREB%       = OREB_Boca / (OREB_Boca + DREB_Rival) × 100
  SC Opport.  = Eventos REBOTE-OFENSIVO de Boca en PBP
  SC PPP      = Puntos marcados en la misma posesión tras el OREB
                 (posesión termina al anotar 2P/3P, pérdida, o DREB rival)
  OREB Total  = Suma de la columna OReb del box score
""")

# ─── EXPORT TO WORD ──────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = '002D62'
GOLD = 'F5C518'
WHITE = 'FFFFFF'
LIGHT_GRAY = 'F2F2F2'
DARK_TEXT = '1A1A1A'

OUT_OREB = '/Users/ramiellero/liga_argentina/Boca_OREB_Analysis.docx'

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_font(cell, text, bold=False, size=10, color=None, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.clear()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)

def add_section_heading(doc, title, subtitle=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    r, g, b = int(BLUE[0:2],16), int(BLUE[2:4],16), int(BLUE[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].italic = True

def add_oreb_table(doc, rows, period_name, n_games):
    """Build a Win/Loss OREB comparison table"""
    cols = [
        ('Escenario',        'Escenario',        1.4, 'left'),
        ('Partidos',         'PJ',               0.5, 'center'),
        ('OREB_total',       'OREB\nTotal',       0.85,'center'),
        ('OREB_pg',          'OREB/G',            0.7, 'center'),
        ('OREB_pct',         'OREB%',             0.7, 'center'),
        ('SC_oportunidades', 'SC\nOpport.',        0.8, 'center'),
        ('SC_opp_pg',        'SC\nOpp/G',          0.7, 'center'),
        ('SC_pts_total',     'SC Pts\nTotal',      0.85,'center'),
        ('SC_pts_pg',        'SC Pts/G',           0.75,'center'),
        ('SC_PPP',           'PPP',               0.65,'center'),
    ]

    n_cols = len(cols)
    table  = doc.add_table(rows=1 + len(rows) + (1 if len(rows)==2 else 0), cols=n_cols)
    table.style = 'Table Grid'

    # Set column widths
    for ci, (_, _, w, _) in enumerate(cols):
        for row in table.rows:
            row.cells[ci].width = Inches(w)

    # Header row
    hdr = table.rows[0]
    for ci, (key, label, w, align) in enumerate(cols):
        set_cell_bg(hdr.cells[ci], BLUE)
        set_cell_font(hdr.cells[ci], label, bold=True, size=8, color=GOLD, align='center')

    # Data rows
    for ri, r in enumerate(rows):
        row  = table.rows[ri + 1]
        bg   = 'E8F5E9' if r['Escenario'] == 'Victoria' else 'FFEBEE'
        for ci, (key, label, w, align) in enumerate(cols):
            set_cell_bg(row.cells[ci], bg)
            val = r[key]
            if key == 'OREB_pct':
                txt = f"{val}%"
            elif key == 'SC_PPP':
                txt = f"{val:.3f}"
            else:
                txt = str(val)
            bold_cell = key in ('Escenario', 'OREB_pg', 'OREB_pct', 'SC_pts_pg', 'SC_PPP')
            set_cell_font(row.cells[ci], txt, bold=bold_cell, size=9,
                          color=DARK_TEXT, align=align)

    # Differential row
    if len(rows) == 2:
        w_row, l_row = rows[0], rows[1]
        diff_row = table.rows[len(rows) + 1]
        set_cell_bg(diff_row.cells[0], '37474F')
        set_cell_font(diff_row.cells[0], 'Δ Vic – Der', bold=True, size=8,
                      color=GOLD, align='left')
        diff_map = {
            'Partidos':         '',
            'OREB_total':       '',
            'OREB_pg':          f"{w_row['OREB_pg'] - l_row['OREB_pg']:+.1f}",
            'OREB_pct':         f"{w_row['OREB_pct'] - l_row['OREB_pct']:+.1f}pp",
            'SC_oportunidades': '',
            'SC_opp_pg':        f"{w_row['SC_opp_pg'] - l_row['SC_opp_pg']:+.1f}",
            'SC_pts_total':     '',
            'SC_pts_pg':        f"{w_row['SC_pts_pg'] - l_row['SC_pts_pg']:+.1f}",
            'SC_PPP':           f"{w_row['SC_PPP'] - l_row['SC_PPP']:+.3f}",
        }
        for ci, (key, label, w, align) in enumerate(cols):
            if ci == 0:
                continue
            set_cell_bg(diff_row.cells[ci], '37474F')
            txt = diff_map.get(key, '')
            set_cell_font(diff_row.cells[ci], txt, bold=True, size=8,
                          color=WHITE, align='center')

    return table


doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BOCA JUNIORS')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x2D, 0x62)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Análisis de Rebotes Ofensivos y Puntos de Segunda Oportunidad')
run2.font.size = Pt(13)
run2.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Liga Nacional Argentina 2025-26  |  14/04/2026')
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# ── GLOSARIO SC ──────────────────────────────────────────────────────────────
p_gl = doc.add_paragraph()
run_gl = p_gl.add_run('Nota sobre SC (Segunda Chance / Segunda Oportunidad):  ')
run_gl.bold = True
run_gl.font.size = Pt(9)
run_gl.font.color.rgb = RGBColor(0x00, 0x2D, 0x62)
run_gl2 = p_gl.add_run(
    'Se denomina "SC" a toda posesión que comienza inmediatamente después de que Boca captura un rebote ofensivo. '
    'La posesión de segunda oportunidad finaliza cuando: (a) Boca convierte un tiro de 2 o 3 puntos, '
    '(b) el rival obtiene un rebote defensivo, o (c) Boca pierde el balón. '
    'Los tiros libres dentro de esa posesión también se suman al cómputo de puntos. '
    'La métrica PPP (Puntos Por Posesión) refleja la eficiencia de conversión en esas posesiones.'
)
run_gl2.font.size = Pt(9)
run_gl2.italic = True

doc.add_paragraph()

# ── DEFINICIONES COLUMNAS ─────────────────────────────────────────────────────
p_def = doc.add_paragraph()
run_d = p_def.add_run('Columnas de la tabla:  ')
run_d.bold = True
run_d.font.size = Pt(9)
defs = [
    ('OREB Total', 'Total de rebotes ofensivos capturados por Boca (fuente: box score).'),
    ('OREB/G',     'Promedio de rebotes ofensivos por partido.'),
    ('OREB%',      'OREB_Boca ÷ (OREB_Boca + DREB_Rival) × 100. Porcentaje de rebotes ofensivos disponibles que captura Boca.'),
    ('SC Opport.', 'Cantidad de posesiones de segunda oportunidad generadas (= número de OREB en PBP).'),
    ('SC Opp/G',   'Posesiones de segunda oportunidad por partido.'),
    ('SC Pts Total','Puntos totales convertidos en posesiones de SC.'),
    ('SC Pts/G',   'Puntos de segunda oportunidad por partido.'),
    ('PPP',        'Puntos por posesión de SC (SC Pts Total ÷ SC Opport.).'),
]
run_d2 = p_def.add_run('\n' + '\n'.join(f'  • {k}: {v}' for k, v in defs))
run_d2.font.size = Pt(8.5)

doc.add_paragraph()

# ── TABLAS POR PERÍODO ────────────────────────────────────────────────────────
period_subtitles = {
    'PRE':  'Período anterior al cambio de cuerpo técnico (antes del 02/02/2026)',
    'CASA': 'Desde la llegada de Casalanguida (desde 10/02/2026)',
    'L5':   'Los últimos 5 partidos disputados',
}

for lbl, name in PERIOD_LABELS.items():
    n_games = len(get_dates(lbl))
    add_section_heading(doc, f'{name}  ({n_games} partidos)', period_subtitles[lbl])
    rows = oreb_analysis(lbl)
    if rows:
        add_oreb_table(doc, rows, name, n_games)
    else:
        doc.add_paragraph('Sin datos para este período.')
    doc.add_paragraph()

# ── INSIGHTS ─────────────────────────────────────────────────────────────────
add_section_heading(doc, 'Insights Clave')

insights = [
    ('El rebote ofensivo no predice victoria en la Era Casalanguida.',
     'Con Casalanguida, Boca gana capturando menos OREB/G (8.0 vs 10.5 en derrotas). '
     'El nuevo sistema prioriza eficiencia en primera posesión sobre segundas oportunidades.'),
    ('Caída pronunciada en SC Pts/G en victorias recientes.',
     'En los últimos 5, Boca convierte solo 6.0 SC Pts/G en victorias vs 15.9 en el período Pre-Casalanguida. '
     'Menor dependencia del cristal ofensivo como fuente de puntos.'),
    ('PPP en SC bajó de 1.39 → 0.82 en victorias (L5).',
     'La eficiencia de conversión en segundas oportunidades recientes es la más baja del año. '
     'Puede indicar posesiones de menor calidad luego del OREB (más tiros forzados).'),
    ('OREB% en victorias cayó 35.1% → 26.6% → 23.4% a lo largo del año.',
     'Tendencia consistente: Boca controla cada vez menos el cristal ofensivo en sus victorias. '
     'Correlaciona con un estilo de juego más dinámico y con menor tiempo en la pintura.'),
    ('Muestra reducida en derrotas Era Casalanguida (n=2).',
     'Los valores de "Derrota" en ese período y en L5 corresponden exactamente a los mismos 2 partidos. '
     'Tratar ese diferencial con cautela; se necesita más muestra para extraer conclusiones definitivas.'),
]

for title, body in insights:
    p_i = doc.add_paragraph(style='List Bullet')
    run_it = p_i.add_run(title + '  ')
    run_it.bold = True
    run_it.font.size = Pt(9.5)
    run_ib = p_i.add_run(body)
    run_ib.font.size = Pt(9.5)

doc.add_paragraph()

# ── NOTA METODOLÓGICA ─────────────────────────────────────────────────────────
p_m = doc.add_paragraph()
run_m = p_m.add_run('Metodología: ')
run_m.bold = True
run_m.font.size = Pt(8.5)
run_m2 = p_m.add_run(
    'OREB totales extraídos del box score oficial. Posesiones de SC reconstruidas evento a evento '
    'desde el play-by-play (evento REBOTE-OFENSIVO de Boca → seguimiento hasta fin de posesión). '
    'Períodos: Pre-Casalanguida < 02/02/2026; Era Casalanguida ≥ 10/02/2026; L5 = últimos 5 partidos.'
)
run_m2.font.size = Pt(8.5)
run_m2.italic = True

doc.save(OUT_OREB)
print(f'\n✓ Word exportado → {OUT_OREB}')

