"""
Boca Juniors – Scouting Report Generator
Liga Nacional Argentina 2025-26
"""

import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings('ignore')

# ─── PATHS ───────────────────────────────────────────────────────────────────
BOX  = '/Users/ramiellero/liga_argentina/liga_argentina/docs/liga_nacional/liga_nacional.csv'
PBP  = '/Users/ramiellero/liga_argentina/liga_argentina/docs/liga_nacional/liga_nacional_pbp.csv'
SHOT = '/Users/ramiellero/liga_argentina/liga_argentina/docs/liga_nacional/liga_nacional_shots.csv'
OUT  = '/Users/ramiellero/liga_argentina/Scouting_Boca_Juniors.docx'

# ─── LOAD DATA ───────────────────────────────────────────────────────────────
print("Loading data...")
box  = pd.read_csv(BOX)
pbp  = pd.read_csv(PBP)
shot = pd.read_csv(SHOT)

box['Fecha_dt']  = pd.to_datetime(box['Fecha'],  format='%d/%m/%Y')
pbp['Fecha_dt']  = pd.to_datetime(pbp['Fecha'],  format='%d/%m/%Y')
shot['Fecha_dt'] = pd.to_datetime(shot['Fecha'], format='%d/%m/%Y')

CUTOFF = pd.Timestamp('2026-02-02')

# Boca games (sorted)
boca_box = box[box['Equipo'] == 'BOCA'].copy()
game_dates = sorted(boca_box[boca_box['Apellido'] == 'TOTALES']['Fecha_dt'].unique())
last5_dates = game_dates[-5:]

def period_label(dt):
    if dt < CUTOFF:     return 'PRE'
    else:               return 'CASA'

def get_period_dates(label):
    if label == 'PRE':   return [d for d in game_dates if d < CUTOFF]
    if label == 'CASA':  return [d for d in game_dates if d >= CUTOFF]
    if label == 'L5':    return last5_dates

print(f"Pre-Casalanguida: {len(get_period_dates('PRE'))} games")
print(f"Era Casalanguida: {len(get_period_dates('CASA'))} games")
print(f"Últimos 5: {len(last5_dates)} games")

# ─── HELPER: filter box by period ────────────────────────────────────────────
def box_period(label, team='BOCA'):
    dates = get_period_dates(label)
    df = box[(box['Equipo'] == team) & (box['Fecha_dt'].isin(dates))]
    return df

def totals_period(label, team='BOCA'):
    return box_period(label, team)[box_period(label, team)['Apellido'] == 'TOTALES']

def players_period(label, team='BOCA'):
    return box_period(label, team)[box_period(label, team)['Apellido'] != 'TOTALES']

def opp_totals(label):
    """Opponent totals for Boca games"""
    dates   = get_period_dates(label)
    boca_games = boca_box[(boca_box['Apellido'] == 'TOTALES') & (boca_box['Fecha_dt'].isin(dates))][['Fecha_dt','IdPartido','Rival']]
    opp_rows = []
    for _, row in boca_games.iterrows():
        opp = box[(box['IdPartido'] == row['IdPartido']) & (box['Equipo'] == row['Rival']) & (box['Apellido'] == 'TOTALES')]
        opp_rows.append(opp)
    return pd.concat(opp_rows) if opp_rows else pd.DataFrame()

# ─── POSSESSION ESTIMATE ─────────────────────────────────────────────────────
def poss(df_totals):
    """Estimate possessions per game from totals rows"""
    fga  = df_totals['T2I'] + df_totals['T3I']
    fta  = df_totals['T1I']
    oreb = df_totals['OReb']
    tov  = df_totals['Perdidas']
    return fga + 0.44 * fta - oreb + tov

# ─── 1. TEAM METRICS ─────────────────────────────────────────────────────────
print("Calculating team metrics...")

def team_metrics(label):
    t  = totals_period(label)
    o  = opp_totals(label)
    n  = len(t)
    if n == 0:
        return {}

    fga   = (t['T2I'] + t['T3I']).sum()
    fgm   = (t['T2A'] + t['T3A']).sum()
    t3a   = t['T3A'].sum()
    t3i   = t['T3I'].sum()
    t2a   = t['T2A'].sum()
    t2i   = t['T2I'].sum()
    fta   = t['T1I'].sum()
    ftm   = t['T1A'].sum()
    pts   = t['Puntos'].sum()
    oreb  = t['OReb'].sum()
    dreb  = t['DReb'].sum()
    tov   = t['Perdidas'].sum()
    ast   = t['Asistencias'].sum()
    mins  = t['Segundos jugados'].sum() / 60

    efg   = (fgm + 0.5 * t3a) / fga if fga > 0 else 0
    t3rate = t3i / fga if fga > 0 else 0
    ftr   = fta / fga if fga > 0 else 0

    # pct shots by zone (mid = 2pt non-paint approximation via %T2 - inside)
    pct_3  = t3i / fga if fga > 0 else 0
    pct_2  = t2i / fga if fga > 0 else 0

    team_poss = (poss(t)).sum()
    opp_poss  = poss(o).sum() if len(o) > 0 else team_poss

    opp_pts   = o['Puntos'].sum() if len(o) > 0 else 0
    opp_oreb  = o['OReb'].sum()   if len(o) > 0 else 0
    opp_dreb  = o['DReb'].sum()   if len(o) > 0 else 0

    ortg  = pts   / team_poss * 100 if team_poss > 0 else 0
    drtg  = opp_pts / opp_poss * 100 if opp_poss > 0 else 0
    net   = ortg - drtg

    # Pace = possessions per 40 min (Argentine games)
    total_secs  = t['Segundos jugados'].sum()   # sum of all player-seconds ÷ 5
    game_mins   = total_secs / 5 / 60
    pace        = (team_poss + opp_poss) / 2 / n if n > 0 else 0

    oreb_pct = oreb / (oreb + opp_dreb) * 100 if (oreb + opp_dreb) > 0 else 0
    dreb_pct = dreb / (dreb + opp_oreb) * 100 if (dreb + opp_oreb) > 0 else 0

    # Per-quarter points
    # Need to calculate from PBP
    dates = get_period_dates(label)
    boca_ids = boca_box[(boca_box['Apellido'] == 'TOTALES') & (boca_box['Fecha_dt'].isin(dates))]['IdPartido'].tolist()
    boca_pbp = pbp[pbp['IdPartido'].isin(boca_ids)].copy()
    boca_pbp2 = boca_pbp[(boca_pbp['Equipo_local'] == 'BOCA') | (boca_pbp['Equipo_visitante'] == 'BOCA')].copy()

    q_pts = {}
    for q in [1,2,3,4]:
        q_data = boca_pbp2[(boca_pbp2['Periodo'] == q) & (boca_pbp2['Tipo'].str.startswith('FINAL-PERIODO', na=False))]
        if len(q_data) == 0:
            q_pts[q] = None
            continue
        vals = []
        for _, row in q_data.iterrows():
            is_local = row['Equipo_local'] == 'BOCA'
            score = row['Marcador_local'] if is_local else row['Marcador_visitante']
            vals.append(score)
        # Q1 pts = score at end of Q1; Q2 pts = score at end of Q2 - Q1; etc.
        q_pts[q] = vals

    # Simpler: accumulate from end-of-period markers
    qtr_pts = {1:[], 2:[], 3:[], 4:[]}
    for gid in boca_ids:
        g = boca_pbp[(boca_pbp['IdPartido'] == gid)].sort_values('NumAccion')
        is_local = g['Equipo_local'].iloc[0] == 'BOCA'
        prev_score = 0
        for q in [1,2,3,4]:
            end = g[(g['Periodo'] == q) & (g['Tipo'].str.contains('FINAL-PERIODO|FINAL-PARTIDO', na=False))]
            if len(end) == 0:
                continue
            row   = end.iloc[-1]
            score = row['Marcador_local'] if is_local else row['Marcador_visitante']
            qtr_pts[q].append(score - prev_score)
            prev_score = score

    return {
        'n': n,
        'pts_pg':  pts / n,
        'efg':     efg * 100,
        't3_rate': t3rate * 100,
        'ftr':     ftr * 100,
        'pct_3':   pct_3 * 100,
        'pct_2':   pct_2 * 100,
        'ortg':    ortg,
        'drtg':    drtg,
        'net':     net,
        'pace':    pace,
        'oreb_pct': oreb_pct,
        'dreb_pct': dreb_pct,
        'q1': np.mean(qtr_pts[1]) if qtr_pts[1] else None,
        'q2': np.mean(qtr_pts[2]) if qtr_pts[2] else None,
        'q3': np.mean(qtr_pts[3]) if qtr_pts[3] else None,
        'q4': np.mean(qtr_pts[4]) if qtr_pts[4] else None,
        'tov_pct': tov / team_poss * 100 if team_poss > 0 else 0,
        'ast_pg': ast / n,
    }

m_pre  = team_metrics('PRE')
m_casa = team_metrics('CASA')
m_l5   = team_metrics('L5')
print("Team metrics done.")

# ─── 2. PLAYER METRICS ───────────────────────────────────────────────────────
print("Calculating player metrics...")

def player_stats(label):
    df = players_period(label)
    df = df[df['Segundos jugados'] > 0].copy()
    n  = len(get_period_dates(label))

    # Team totals for USG
    t   = totals_period(label)
    tm_fga  = (t['T2I'] + t['T3I']).sum()
    tm_fta  = t['T1I'].sum()
    tm_tov  = t['Perdidas'].sum()
    tm_fgm  = (t['T2A'] + t['T3A']).sum()
    tm_ast  = t['Asistencias'].sum()

    rows = []
    for name, grp in df.groupby('Nombre completo'):
        mins  = grp['Segundos jugados'].sum() / 60
        gp    = grp['Fecha'].nunique()
        fga   = (grp['T2I'] + grp['T3I']).sum()
        fgm   = (grp['T2A'] + grp['T3A']).sum()
        t3a   = grp['T3A'].sum()
        t3i   = grp['T3I'].sum()
        fta   = grp['T1I'].sum()
        ftm   = grp['T1A'].sum()
        pts   = grp['Puntos'].sum()
        ast   = grp['Asistencias'].sum()
        tov   = grp['Perdidas'].sum()

        efg   = (fgm + 0.5*t3a) / fga * 100 if fga > 0 else None
        usg   = (fga + 0.44*fta + tov) / (tm_fga + 0.44*tm_fta + tm_tov) * 100 * 5 if (tm_fga+tm_fta+tm_tov) > 0 else None
        # Scale USG by minutes ratio
        if n > 0 and mins > 0:
            pass  # simple version without minutes scaling for now

        t3rate = t3i / fga * 100 if fga > 0 else None
        ast_pct = ast / tm_fgm * 100 if tm_fgm > 0 else None
        tov_pct = tov / (fga + 0.44*fta + tov) * 100 if (fga + 0.44*fta + tov) > 0 else None

        rows.append({
            'Jugador': name,
            'GP': gp,
            'Min_total': mins,
            'Min_pg': mins / gp if gp > 0 else 0,
            'Pts_pg': pts / gp if gp > 0 else 0,
            'eFG%': efg,
            'USG%': usg,
            '3P_rate': t3rate,
            'AST%': ast_pct,
            'TOV%': tov_pct,
        })
    return pd.DataFrame(rows).sort_values('Min_pg', ascending=False)

ps_pre  = player_stats('PRE')
ps_casa = player_stats('CASA')
ps_l5   = player_stats('L5')
ps_all  = player_stats.__wrapped__ if hasattr(player_stats,'__wrapped__') else None

# Full season
def player_stats_all():
    df = boca_box[boca_box['Apellido'] != 'TOTALES'].copy()
    df = df[df['Segundos jugados'] > 0]
    t  = boca_box[boca_box['Apellido'] == 'TOTALES']
    tm_fga = (t['T2I'] + t['T3I']).sum()
    tm_fta = t['T1I'].sum()
    tm_tov = t['Perdidas'].sum()
    tm_fgm = (t['T2A'] + t['T3A']).sum()
    rows = []
    for name, grp in df.groupby('Nombre completo'):
        mins = grp['Segundos jugados'].sum() / 60
        gp   = grp['Fecha'].nunique()
        fga  = (grp['T2I'] + grp['T3I']).sum()
        fgm  = (grp['T2A'] + grp['T3A']).sum()
        t3a  = grp['T3A'].sum()
        t3i  = grp['T3I'].sum()
        fta  = grp['T1I'].sum()
        pts  = grp['Puntos'].sum()
        ast  = grp['Asistencias'].sum()
        tov  = grp['Perdidas'].sum()
        efg   = (fgm + 0.5*t3a) / fga * 100 if fga > 0 else None
        usg   = (fga + 0.44*fta + tov) / (tm_fga + 0.44*tm_fta + tm_tov) * 100 * 5 if (tm_fga+tm_fta+tm_tov) > 0 else None
        t3rate = t3i / fga * 100 if fga > 0 else None
        ast_pct = ast / tm_fgm * 100 if tm_fgm > 0 else None
        tov_pct = tov / (fga + 0.44*fta + tov) * 100 if (fga + 0.44*fta + tov) > 0 else None
        rows.append({
            'Jugador': name, 'GP': gp,
            'Min_pg': mins/gp if gp>0 else 0,
            'Pts_pg': pts/gp if gp>0 else 0,
            'eFG%': efg, 'USG%': usg, '3P_rate': t3rate,
            'AST%': ast_pct, 'TOV%': tov_pct,
        })
    return pd.DataFrame(rows).sort_values('Min_pg', ascending=False)

ps_full = player_stats_all()

# Top 8 by minutes in last 5
top8 = ps_l5.head(8)['Jugador'].tolist()
print(f"Top 8 by L5 minutes: {top8}")

# ─── 3. SHOT DISTRIBUTION ────────────────────────────────────────────────────
print("Calculating shot distribution...")

ZONE_MAP = {
    'Z1':'Aro/Pintura IZ', 'Z2':'Aro/Pintura CE', 'Z3':'Aro/Pintura DE',
    'Z4':'Pintura IZ', 'Z5':'Pintura DE',
    'Z6':'Media Dist IZ', 'Z7':'Media Dist CE-IZ', 'Z8':'Media Dist CE',
    'Z9':'Media Dist CE-DE', 'Z10':'Media Dist DE',
    'Z11':'Triple IZ', 'Z12':'Triple CE-IZ', 'Z13':'Triple CE',
    'Z14':'Triple CE-DE', 'Z15':'Triple DE',
}

def shot_zone_base(z):
    if z.startswith('Z1') or z.startswith('Z2') or z.startswith('Z3') or z.startswith('Z4') or z.startswith('Z5'):
        return 'Aro/Pintura'
    elif z.startswith('Z1') or z.startswith('Z6') or z.startswith('Z7') or z.startswith('Z8') or z.startswith('Z9') or z.startswith('Z10'):
        return 'Media Distancia'
    else:
        return 'Triple'

def categorize_zone(z):
    if pd.isna(z): return 'Desconocida'
    z = str(z).split('-')[0]  # get Z1, Z2, etc.
    num = int(z[1:]) if z[1:].isdigit() else 0
    if num <= 5:   return 'Aro/Pintura'
    elif num <= 10: return 'Media Distancia'
    else:          return 'Triple'

def shot_dist(label):
    dates = get_period_dates(label)
    boca_ids = boca_box[(boca_box['Apellido'] == 'TOTALES') & (boca_box['Fecha_dt'].isin(dates))]['IdPartido'].tolist()
    boca_shot = shot[(shot['IdPartido'].isin(boca_ids)) & (shot['Equipo'] == 'BOCA')].copy()
    boca_shot['cat'] = boca_shot['Zona'].apply(categorize_zone)
    total = len(boca_shot)
    result = {}
    for cat in ['Aro/Pintura', 'Media Distancia', 'Triple']:
        sub = boca_shot[boca_shot['cat'] == cat]
        freq = len(sub) / total * 100 if total > 0 else 0
        conv = sub[sub['Resultado'] == 'CONVERTIDO']
        # eFG for zone
        t3 = sub[sub['Tipo'] == 'TIRO3']
        t2 = sub[sub['Tipo'] == 'TIRO2']
        fgm = len(sub[sub['Resultado'] == 'CONVERTIDO'])
        t3m = len(t3[t3['Resultado'] == 'CONVERTIDO'])
        fga = len(sub)
        efg_z = (fgm + 0.5*t3m) / fga * 100 if fga > 0 else None
        fg_pct = fgm / fga * 100 if fga > 0 else None
        result[cat] = {'freq': freq, 'FG%': fg_pct, 'eFG%': efg_z, 'FGA': fga}
    return result, total

sd_pre,  n_pre  = shot_dist('PRE')
sd_casa, n_casa = shot_dist('CASA')
sd_l5,   n_l5   = shot_dist('L5')
print("Shot distribution done.")

# ─── 4. CLUTCH ───────────────────────────────────────────────────────────────
print("Calculating clutch metrics...")

def parse_time(t):
    """Parse mm:ss -> seconds remaining"""
    try:
        parts = str(t).split(':')
        return int(parts[0])*60 + int(parts[1])
    except:
        return None

def clutch_stats(label):
    dates = get_period_dates(label)
    boca_ids = boca_box[(boca_box['Apellido'] == 'TOTALES') & (boca_box['Fecha_dt'].isin(dates))]['IdPartido'].tolist()
    g = pbp[pbp['IdPartido'].isin(boca_ids)].copy()

    clutch_shots = []
    for gid in boca_ids:
        game = g[g['IdPartido'] == gid].sort_values('NumAccion')
        is_local = game['Equipo_local'].iloc[0] == 'BOCA'
        for _, row in game.iterrows():
            if row['Periodo'] < 4: continue
            secs = parse_time(row['Tiempo'])
            if secs is None: continue
            if row['Periodo'] == 4 and secs > 300: continue  # only last 5 min
            # Check score diff
            diff = abs(row['Marcador_local'] - row['Marcador_visitante'])
            if diff > 5: continue
            # Is it a Boca shot?
            boca_side = 'LOCAL' if is_local else 'VISITANTE'
            if row['Equipo_lado'] != boca_side: continue
            # Only field goals (exclude free throws)
            tipo = str(row['Tipo'])
            if tipo not in ('CANASTA-2P', 'CANASTA-3P', 'TIRO2-FALLADO', 'TIRO3-FALLADO'):
                continue
            made = tipo.startswith('CANASTA')
            t3   = '3P' in tipo
            pts_val = 3 if tipo == 'CANASTA-3P' else (2 if tipo == 'CANASTA-2P' else 0)
            clutch_shots.append({
                'Jugador': row['Jugador'],
                'made': made,
                'pts': pts_val,
                't3': t3,
            })

    if not clutch_shots:
        return pd.DataFrame()

    df_c = pd.DataFrame(clutch_shots)
    df_c['fga'] = 1
    df_c['fgm'] = df_c['made'].astype(int)
    df_c['t3a'] = df_c['t3'].astype(int)
    df_c['t3m'] = ((df_c['t3']) & (df_c['made'])).astype(int)

    out = df_c.groupby('Jugador').agg(
        FGA=('fga','sum'), FGM=('fgm','sum'),
        T3A=('t3a','sum'), T3M=('t3m','sum'),
    ).reset_index()
    out['eFG%'] = (out['FGM'] + 0.5*out['T3M']) / out['FGA'] * 100
    out['USG_clutch'] = out['FGA'] / out['FGA'].sum() * 100
    return out.sort_values('FGA', ascending=False)

cl_pre  = clutch_stats('PRE')
cl_casa = clutch_stats('CASA')
cl_l5   = clutch_stats('L5')
print("Clutch done.")

# ─── 5. CONNECTIONS (PLAYMAKING) ─────────────────────────────────────────────
print("Calculating connections...")

def connections(label):
    dates = get_period_dates(label)
    boca_ids = boca_box[(boca_box['Apellido'] == 'TOTALES') & (boca_box['Fecha_dt'].isin(dates))]['IdPartido'].tolist()
    g = pbp[pbp['IdPartido'].isin(boca_ids)].copy()

    pairs = []
    for gid in boca_ids:
        game = g[g['IdPartido'] == gid].sort_values('NumAccion').reset_index(drop=True)
        is_local = game['Equipo_local'].iloc[0] == 'BOCA'
        boca_side = 'LOCAL' if is_local else 'VISITANTE'
        for i, row in game.iterrows():
            if row['Equipo_lado'] != boca_side: continue
            if str(row['Tipo']) != 'ASISTENCIA': continue
            assist_player = row['Jugador']
            # Look for the basket just before this assist
            for j in range(i-1, max(i-5,-1), -1):
                prev = game.iloc[j]
                if prev['Equipo_lado'] != boca_side: continue
                if str(prev['Tipo']).startswith('CANASTA'):
                    scorer = prev['Jugador']
                    pairs.append((assist_player, scorer))
                    break

    if not pairs:
        return pd.DataFrame()
    df_p = pd.DataFrame(pairs, columns=['Pasador','Anotador'])
    return df_p.groupby(['Pasador','Anotador']).size().reset_index(name='N').sort_values('N', ascending=False).head(5)

cn_pre  = connections('PRE')
cn_casa = connections('CASA')
cn_l5   = connections('L5')
print("Connections done.")

# ─── 6. LINEUPS ──────────────────────────────────────────────────────────────
print("Calculating lineups...")

def extract_lineups(label):
    dates = get_period_dates(label)
    boca_ids = boca_box[(boca_box['Apellido'] == 'TOTALES') & (boca_box['Fecha_dt'].isin(dates))]['IdPartido'].tolist()
    n_games = len(boca_ids)

    lineup_stats = {}

    for gid in boca_ids:
        game_pbp = pbp[pbp['IdPartido'] == gid].sort_values('NumAccion').reset_index(drop=True)
        if len(game_pbp) == 0: continue

        is_local = game_pbp['Equipo_local'].iloc[0] == 'BOCA'
        boca_side = 'LOCAL' if is_local else 'VISITANTE'

        # Track current lineup
        current_lineup = set()
        # Initialize with starters (first CAMBIO-JUGADOR-ENTRA in period 1)
        period1 = game_pbp[game_pbp['Periodo'] == 1]
        for _, row in period1.iterrows():
            if row['Equipo_lado'] != boca_side: continue
            if str(row['Tipo']) == 'CAMBIO-JUGADOR-ENTRA':
                current_lineup.add(str(row['Jugador']))
            if len(current_lineup) == 5:
                break

        segment_start_idx = 0
        segment_start_score_boca = 0
        segment_start_score_opp  = 0

        def get_scores(row):
            if is_local:
                return row['Marcador_local'], row['Marcador_visitante']
            else:
                return row['Marcador_visitante'], row['Marcador_local']

        def get_time_secs(row):
            period = row['Periodo']
            t = parse_time(row['Tiempo'])
            if t is None: return None
            elapsed = (period - 1) * 600 + (600 - t)
            return elapsed

        def record_segment(lineup, start_i, end_i):
            if len(lineup) != 5: return
            if start_i >= end_i: return
            start_row = game_pbp.iloc[start_i]
            end_row   = game_pbp.iloc[end_i]
            t_start = get_time_secs(start_row)
            t_end   = get_time_secs(end_row)
            if t_start is None or t_end is None: return
            seg_mins = (t_end - t_start) / 60
            if seg_mins <= 0: return

            # Count stats in segment
            seg = game_pbp.iloc[start_i:end_i+1]
            boca_seg = seg[seg['Equipo_lado'] == boca_side]
            opp_seg  = seg[seg['Equipo_lado'] != boca_side]

            def count_poss(rows):
                fga = rows['Tipo'].str.startswith('CANASTA', na=False).sum() + \
                      rows['Tipo'].str.startswith('TIRO-FALLADO', na=False).sum()
                oreb = (rows['Tipo'] == 'REBOTE-OFENSIVO').sum()
                tov  = (rows['Tipo'] == 'PERDIDA').sum()
                fta_ev = rows['Tipo'].str.contains('FALTA-RECIBIDA', na=False).sum()
                return fga - oreb + tov + 0.44 * fta_ev

            b_poss = count_poss(boca_seg)
            o_poss = count_poss(opp_seg)

            # Points in segment
            b_pts = 0
            o_pts = 0
            for _, r in seg.iterrows():
                if str(r['Tipo']).startswith('CANASTA'):
                    pts_v = 3 if '3P' in str(r['Tipo']) else (2 if '2P' in str(r['Tipo']) else 1)
                    if r['Equipo_lado'] == boca_side:
                        b_pts += pts_v
                    else:
                        o_pts += pts_v

            boca_3pa = boca_seg['Tipo'].str.contains('3P', na=False).sum()
            boca_fga_total = boca_seg['Tipo'].str.startswith('CANASTA', na=False).sum() + \
                             boca_seg['Tipo'].str.startswith('TIRO-FALLADO', na=False).sum()
            boca_ast = (boca_seg['Tipo'] == 'ASISTENCIA').sum()
            boca_fgm = boca_seg['Tipo'].str.startswith('CANASTA', na=False).sum()
            boca_tov = (boca_seg['Tipo'] == 'PERDIDA').sum()
            boca_fta = boca_seg['Tipo'].str.contains('1P', na=False).sum()

            key = tuple(sorted(lineup))
            if key not in lineup_stats:
                lineup_stats[key] = {
                    'mins':0,'b_poss':0,'o_poss':0,'b_pts':0,'o_pts':0,
                    'b_3pa':0,'b_fga':0,'b_ast':0,'b_fgm':0,'b_tov':0,'b_fta':0,
                    'games': set()
                }
            lineup_stats[key]['mins']   += seg_mins
            lineup_stats[key]['b_poss'] += b_poss
            lineup_stats[key]['o_poss'] += o_poss
            lineup_stats[key]['b_pts']  += b_pts
            lineup_stats[key]['o_pts']  += o_pts
            lineup_stats[key]['b_3pa']  += boca_3pa
            lineup_stats[key]['b_fga']  += boca_fga_total
            lineup_stats[key]['b_ast']  += boca_ast
            lineup_stats[key]['b_fgm']  += boca_fgm
            lineup_stats[key]['b_tov']  += boca_tov
            lineup_stats[key]['b_fta']  += boca_fta
            lineup_stats[key]['games'].add(gid)

        # Process substitutions
        for i, row in game_pbp.iterrows():
            if row['Equipo_lado'] != boca_side: continue
            if str(row['Tipo']) in ['CAMBIO-JUGADOR-ENTRA', 'CAMBIO-JUGADOR-SALE']:
                # Record segment up to this point
                record_segment(current_lineup, segment_start_idx, i)
                segment_start_idx = i
                if str(row['Tipo']) == 'CAMBIO-JUGADOR-ENTRA':
                    current_lineup.add(str(row['Jugador']))
                else:
                    current_lineup.discard(str(row['Jugador']))
            elif str(row['Tipo']) in ['INICIO-PERIODO', 'FINAL-PERIODO', 'INICIO-PARTIDO', 'FINAL-PARTIDO']:
                record_segment(current_lineup, segment_start_idx, i)
                segment_start_idx = i

        # Final segment
        record_segment(current_lineup, segment_start_idx, len(game_pbp)-1)

    # Build dataframe
    rows = []
    for lineup, stats in lineup_stats.items():
        if stats['mins'] < 3: continue  # min filter
        b_p = stats['b_poss']
        o_p = stats['o_poss']
        ortg = stats['b_pts'] / b_p * 100 if b_p > 0 else None
        drtg = stats['o_pts'] / o_p * 100 if o_p > 0 else None
        net  = ortg - drtg if (ortg and drtg) else None
        t3rate = stats['b_3pa'] / stats['b_fga'] * 100 if stats['b_fga'] > 0 else None
        ftr  = stats['b_fta'] / stats['b_fga'] * 100 if stats['b_fga'] > 0 else None
        ast_pct = stats['b_ast'] / stats['b_fgm'] * 100 if stats['b_fgm'] > 0 else None
        tov_pct = stats['b_tov'] / b_p * 100 if b_p > 0 else None
        poss_pg = b_p / len(stats['games']) if stats['games'] else 0
        rows.append({
            'Quinteto': ' / '.join([p.split(',')[0].strip() for p in lineup]),
            'Mins': round(stats['mins'],1),
            'Poss': round(b_p,0),
            'Poss/G': round(poss_pg,1),
            'ORtg': round(ortg,1) if ortg else None,
            'DRtg': round(drtg,1) if drtg else None,
            'Net': round(net,1) if net else None,
            'TOV%': round(tov_pct,1) if tov_pct else None,
            '3PA%': round(t3rate,1) if t3rate else None,
            'FTr': round(ftr,1) if ftr else None,
            'AST%': round(ast_pct,1) if ast_pct else None,
        })

    return pd.DataFrame(rows)

lu_pre  = extract_lineups('PRE')
lu_casa = extract_lineups('CASA')
lu_l5   = extract_lineups('L5')

def top_lineups(df, by='Net', n=5, min_mins=5):
    if df.empty: return df
    df2 = df[df['Mins'] >= min_mins].dropna(subset=[by])
    return df2.nlargest(n, by)

def most_used(df, n=5, min_mins=5):
    if df.empty: return df
    return df[df['Mins'] >= min_mins].nlargest(n, 'Mins')

print("Lineups done.")

# ─── DOCX GENERATION ─────────────────────────────────────────────────────────
print("Generating DOCX...")

BLUE  = RGBColor(0x00, 0x2D, 0x62)   # Boca dark blue
GOLD  = RGBColor(0xF5, 0xC5, 0x18)   # Boca gold
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF2, 0xF2, 0xF2)
DGRAY = RGBColor(0x40, 0x40, 0x40)

def set_cell_bg(cell, color_rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(*color_rgb)
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9, color=None, align='center'):
    cell.text = ''
    p  = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align=='center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text) if text is not None else '—')
    run.bold  = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def add_header_row(table, headers, bg=BLUE, txt_color=WHITE, size=8):
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i < len(row.cells):
            set_cell_bg(row.cells[i], bg)
            cell_text(row.cells[i], h, bold=True, size=size, color=txt_color)

def alt_row(table, row_idx):
    if row_idx % 2 == 0:
        for cell in table.rows[row_idx].cells:
            set_cell_bg(cell, LGRAY)

def add_section(doc, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    # underline
    p2 = doc.add_paragraph()
    run2 = p2.add_run('─' * 80)
    run2.font.size = Pt(7)
    run2.font.color.rgb = GOLD

def add_subsection(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DGRAY

def fmt(val, dec=1, pct=False):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return '—'
    if pct:
        return f"{val:.{dec}f}%"
    return f"{val:.{dec}f}"

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(0.75)
section.right_margin  = Inches(0.75)
section.top_margin    = Inches(0.75)
section.bottom_margin = Inches(0.75)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── TITLE PAGE ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SCOUTING REPORT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('BOCA JUNIORS')
run2.bold = True
run2.font.size = Pt(28)
run2.font.color.rgb = GOLD

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Liga Nacional de Básquet Argentina — Temporada 2025-26')
run3.font.size = Pt(11)
run3.font.color.rgb = DGRAY

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(f'Generado: {datetime.today().strftime("%d/%m/%Y")}')
run4.font.size = Pt(9)
run4.font.color.rgb = DGRAY

doc.add_paragraph()

# Period summary table
pt = doc.add_table(rows=4, cols=4)
pt.style = 'Table Grid'
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_pt = ['Período', 'Fechas', 'Partidos', 'Balance']
add_header_row(pt, headers_pt)
pre_wins  = totals_period('PRE')['Ganado'].sum()
pre_n     = len(totals_period('PRE'))
casa_wins = totals_period('CASA')['Ganado'].sum()
casa_n    = len(totals_period('CASA'))
l5_wins   = totals_period('L5')['Ganado'].sum()
l5_n      = len(totals_period('L5'))

rows_data = [
    ['Pre-Casalanguida',  '24/09/2025 – 30/01/2026', str(pre_n),  f"{int(pre_wins)}-{pre_n-int(pre_wins)}"],
    ['Era Casalanguida',  '10/02/2026 – 11/04/2026', str(casa_n), f"{int(casa_wins)}-{casa_n-int(casa_wins)}"],
    ['Últimos 5',         '25/03/2026 – 11/04/2026', '5',          f"{int(l5_wins)}-{5-int(l5_wins)}"],
]
for i, rd in enumerate(rows_data):
    for j, val in enumerate(rd):
        cell_text(pt.rows[i+1].cells[j], val, size=9)
    alt_row(pt, i+1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: TEAM PROFILE
# ══════════════════════════════════════════════════════════════════════════════
add_section(doc, '1. PERFIL DE TIRO — EVOLUCIÓN')

add_subsection(doc, '1.1 Distribución y Eficiencia de Tiro')
t_shot = doc.add_table(rows=5, cols=8)
t_shot.style = 'Table Grid'
t_shot.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['Período','% T2 (FGA)','% T3 (FGA)','3P Rate','FT Rate','T2%','T3%','eFG%']
add_header_row(t_shot, hdrs)

def get_shot_row(label, name):
    t  = totals_period(label)
    fga = (t['T2I'] + t['T3I']).sum()
    fgm = (t['T2A'] + t['T3A']).sum()
    t3a = t['T3A'].sum()
    t3i = t['T3I'].sum()
    t2a = t['T2A'].sum()
    t2i = t['T2I'].sum()
    fta = t['T1I'].sum()
    efg = (fgm + 0.5*t3a) / fga * 100 if fga > 0 else None
    return [
        name,
        fmt(t2i/fga*100 if fga>0 else None),
        fmt(t3i/fga*100 if fga>0 else None),
        fmt(t3i/fga*100 if fga>0 else None),
        fmt(fta/fga*100 if fga>0 else None),
        fmt(t2a/t2i*100 if t2i>0 else None),
        fmt(t3a/t3i*100 if t3i>0 else None),
        fmt(efg),
    ]

for i, (lbl, nm) in enumerate([('PRE','Pre-Casalanguida'),('CASA','Era Casalanguida'),('L5','Últimos 5')]):
    row_vals = get_shot_row(lbl, nm)
    for j, v in enumerate(row_vals):
        cell_text(t_shot.rows[i+1].cells[j], v, size=9,
                  bold=(j==0), color=(BLUE if j==0 else None))
    alt_row(t_shot, i+1)

# ── RITMO Y EFICIENCIA ───────────────────────────────────────────────────────
add_section(doc, '2. RITMO Y EFICIENCIA')

t_eff = doc.add_table(rows=4, cols=7)
t_eff.style = 'Table Grid'
t_eff.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs2 = ['Período', 'Pace', 'Pts/G', 'ORtg', 'DRtg', 'Net Rtg', 'TOV%']
add_header_row(t_eff, hdrs2)

for i, (m, nm) in enumerate([(m_pre,'Pre-Casalanguida'),(m_casa,'Era Casalanguida'),(m_l5,'Últimos 5')]):
    vals = [
        nm,
        fmt(m.get('pace')),
        fmt(m.get('pts_pg')),
        fmt(m.get('ortg')),
        fmt(m.get('drtg')),
        fmt(m.get('net')),
        fmt(m.get('tov_pct')),
    ]
    for j, v in enumerate(vals):
        net_val = m.get('net')
        net_color = None
        if j == 5 and net_val is not None:
            net_color = RGBColor(0,128,0) if net_val >= 0 else RGBColor(180,0,0)
        cell_text(t_eff.rows[i+1].cells[j], v, size=9,
                  bold=(j==0 or j==5),
                  color=(BLUE if j==0 else net_color))
    alt_row(t_eff, i+1)

# ── REBOTE ──────────────────────────────────────────────────────────────────
add_section(doc, '3. REBOTE')

t_reb = doc.add_table(rows=4, cols=3)
t_reb.style = 'Table Grid'
t_reb.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(t_reb, ['Período', '%OREB', '%DREB'])

for i, (m, nm) in enumerate([(m_pre,'Pre-Casalanguida'),(m_casa,'Era Casalanguida'),(m_l5,'Últimos 5')]):
    vals = [nm, fmt(m.get('oreb_pct')), fmt(m.get('dreb_pct'))]
    for j, v in enumerate(vals):
        cell_text(t_reb.rows[i+1].cells[j], v, size=9, bold=(j==0), color=(BLUE if j==0 else None))
    alt_row(t_reb, i+1)

# ── PUNTOS POR CUARTO ────────────────────────────────────────────────────────
add_section(doc, '4. PUNTOS POR CUARTO')

t_qtr = doc.add_table(rows=4, cols=5)
t_qtr.style = 'Table Grid'
t_qtr.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(t_qtr, ['Período', 'Q1', 'Q2', 'Q3', 'Q4'])

for i, (m, nm) in enumerate([(m_pre,'Pre-Casalanguida'),(m_casa,'Era Casalanguida'),(m_l5,'Últimos 5')]):
    vals = [nm, fmt(m.get('q1')), fmt(m.get('q2')), fmt(m.get('q3')), fmt(m.get('q4'))]
    for j, v in enumerate(vals):
        cell_text(t_qtr.rows[i+1].cells[j], v, size=9, bold=(j==0), color=(BLUE if j==0 else None))
    alt_row(t_qtr, i+1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: PLAYERS
# ══════════════════════════════════════════════════════════════════════════════
add_section(doc, '5. ANÁLISIS DE JUGADORES')
add_subsection(doc, f'Top 8 por minutos en Últimos 5: {", ".join([p.split(",")[0].strip() for p in top8])}')
doc.add_paragraph()

player_hdrs = ['Jugador', 'GP', 'Min/G', 'Pts/G', 'USG%', 'eFG%', '3P Rate', 'AST%', 'TOV%']

def build_player_table(title, ps_df, players):
    add_subsection(doc, title)
    t = doc.add_table(rows=len(players)+1, cols=len(player_hdrs))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, player_hdrs, size=8)
    for i, player in enumerate(players):
        row_data = ps_df[ps_df['Jugador'] == player]
        if row_data.empty:
            vals = [player.split(',')[0].strip()] + ['—'] * (len(player_hdrs)-1)
        else:
            r = row_data.iloc[0]
            vals = [
                player.split(',')[0].strip(),
                str(int(r['GP'])),
                fmt(r['Min_pg']),
                fmt(r['Pts_pg']),
                fmt(r.get('USG%')),
                fmt(r.get('eFG%')),
                fmt(r.get('3P_rate')),
                fmt(r.get('AST%')),
                fmt(r.get('TOV%')),
            ]
        for j, v in enumerate(vals):
            cell_text(t.rows[i+1].cells[j], v, size=8,
                      bold=(j==0), color=(BLUE if j==0 else None),
                      align='left' if j==0 else 'center')
        alt_row(t, i+1)
    doc.add_paragraph()

build_player_table('5.1 Temporada Completa',    ps_full,  top8)
build_player_table('5.2 Pre-Casalanguida',       ps_pre,   top8)
build_player_table('5.3 Era Casalanguida',        ps_casa,  top8)
build_player_table('5.4 Últimos 5 Partidos',      ps_l5,    top8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: SHOT DISTRIBUTION
# ══════════════════════════════════════════════════════════════════════════════
add_section(doc, '6. DISTRIBUCIÓN DE TIRO POR ZONA')

t_sd = doc.add_table(rows=10, cols=6)
t_sd.style = 'Table Grid'
t_sd.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(t_sd, ['Período', 'Zona', 'FGA', 'Freq%', 'FG%', 'eFG%'])

row_idx = 1
for label, sd, lbl_name in [('PRE',sd_pre,'Pre-Casalanguida'),('CASA',sd_casa,'Era Casalanguida'),('L5',sd_l5,'Últimos 5')]:
    first = True
    for cat in ['Aro/Pintura','Media Distancia','Triple']:
        stats = sd.get(cat, {})
        period_label_cell = lbl_name if first else ''
        first = False
        vals = [
            period_label_cell,
            cat,
            str(stats.get('FGA',0)),
            fmt(stats.get('freq')),
            fmt(stats.get('FG%')),
            fmt(stats.get('eFG%')),
        ]
        for j, v in enumerate(vals):
            cell_text(t_sd.rows[row_idx].cells[j], v, size=8,
                      bold=(j==0 and period_label_cell!=''),
                      color=(BLUE if j==0 else None))
        if period_label_cell:
            alt_row(t_sd, row_idx)
        row_idx += 1

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: CLUTCH
# ══════════════════════════════════════════════════════════════════════════════
add_section(doc, '7. CLUTCH (Últ. 5 min Q4 o Prórroga, Diferencia ≤ 5 pts)')

def build_clutch_table(title, cl_df):
    add_subsection(doc, title)
    if cl_df.empty:
        doc.add_paragraph('Sin datos suficientes.')
        return
    t = doc.add_table(rows=min(len(cl_df),8)+1, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, ['Jugador','FGA','FGM','eFG%','USG% Clutch'])
    for i, (_, row) in enumerate(cl_df.head(8).iterrows()):
        vals = [
            str(row['Jugador']).split(',')[0].strip(),
            str(int(row['FGA'])),
            str(int(row['FGM'])),
            fmt(row['eFG%']),
            fmt(row['USG_clutch']),
        ]
        for j, v in enumerate(vals):
            cell_text(t.rows[i+1].cells[j], v, size=8,
                      bold=(j==0), color=(BLUE if j==0 else None),
                      align='left' if j==0 else 'center')
        alt_row(t, i+1)
    doc.add_paragraph()

build_clutch_table('7.1 Pre-Casalanguida',    cl_pre)
build_clutch_table('7.2 Era Casalanguida',     cl_casa)
build_clutch_table('7.3 Últimos 5 Partidos',   cl_l5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: CONNECTIONS
# ══════════════════════════════════════════════════════════════════════════════
add_section(doc, '8. CONEXIONES (PLAYMAKING) — Top 5 Duplas por Período')

def build_conn_table(title, cn_df):
    add_subsection(doc, title)
    if cn_df.empty:
        doc.add_paragraph('Sin datos.')
        return
    t = doc.add_table(rows=min(len(cn_df),5)+1, cols=3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(t, ['Pasador','Anotador','N'])
    for i, (_, row) in enumerate(cn_df.head(5).iterrows()):
        vals = [
            str(row['Pasador']).split(',')[0].strip(),
            str(row['Anotador']).split(',')[0].strip(),
            str(int(row['N'])),
        ]
        for j, v in enumerate(vals):
            cell_text(t.rows[i+1].cells[j], v, size=9,
                      bold=(j<2), color=(BLUE if j<2 else None))
        alt_row(t, i+1)
    doc.add_paragraph()

build_conn_table('8.1 Pre-Casalanguida',   cn_pre)
build_conn_table('8.2 Era Casalanguida',    cn_casa)
build_conn_table('8.3 Últimos 5 Partidos',  cn_l5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9: LINEUPS
# ══════════════════════════════════════════════════════════════════════════════
add_section(doc, '9. QUINTETOS (LINEUPS) — ANÁLISIS AVANZADO')

lineup_hdrs = ['Quinteto', 'Min', 'Poss', 'P/G', 'ORtg', 'DRtg', 'Net', 'TOV%', '3PA%', 'FTr', 'AST%']

def build_lineup_table(title, lu_df, by='Net'):
    add_subsection(doc, title)
    if lu_df.empty:
        doc.add_paragraph('Sin datos suficientes.')
        return
    top_eff  = top_lineups(lu_df, by='Net', n=5, min_mins=5)
    top_used = most_used(lu_df, n=5, min_mins=5)

    for sub_title, sub_df in [('Más Eficientes (Net Rtg)', top_eff), ('Más Utilizados (Minutos)', top_used)]:
        doc.add_paragraph()
        p = doc.add_paragraph(f'  ► {sub_title}')
        p.runs[0].font.size = Pt(9)
        p.runs[0].bold = True
        if sub_df.empty:
            doc.add_paragraph('  Sin datos.')
            continue
        t = doc.add_table(rows=len(sub_df)+1, cols=len(lineup_hdrs))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(t, lineup_hdrs, size=7)
        for i, (_, row) in enumerate(sub_df.iterrows()):
            net_v = row.get('Net')
            net_color = None
            if net_v is not None and not (isinstance(net_v, float) and np.isnan(net_v)):
                net_color = RGBColor(0,128,0) if float(net_v)>=0 else RGBColor(180,0,0)
            vals = [
                row['Quinteto'],
                fmt(row['Mins']),
                fmt(row['Poss'],0),
                fmt(row['Poss/G']),
                fmt(row['ORtg']),
                fmt(row['DRtg']),
                fmt(row['Net']) if net_v is not None else '—',
                fmt(row['TOV%']),
                fmt(row['3PA%']),
                fmt(row['FTr']),
                fmt(row['AST%']),
            ]
            for j, v in enumerate(vals):
                is_net = (j == 6)
                cell_text(t.rows[i+1].cells[j], v, size=7,
                          bold=(j==0 or is_net),
                          color=(net_color if is_net else None),
                          align='left' if j==0 else 'center')
            alt_row(t, i+1)
    doc.add_paragraph()

build_lineup_table('9.1 Pre-Casalanguida',   lu_pre)
build_lineup_table('9.2 Era Casalanguida',    lu_casa)
build_lineup_table('9.3 Últimos 5 Partidos',  lu_l5)

# ─── SAVE ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\n✓ Saved to: {OUT}")
